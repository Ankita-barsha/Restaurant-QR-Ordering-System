import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_documentation():
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Fonts
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Colors
    NAVY = RGBColor(0x1E, 0x29, 0x3B)
    GOLD = RGBColor(0xD9, 0x77, 0x06)
    ORANGE = RGBColor(0xF9, 0x73, 0x16)
    DARK_GRAY = RGBColor(0x33, 0x41, 0x55)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_after = Pt(4)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = GOLD
        p.paragraph_format.space_after = Pt(4)

    def add_brand_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = ORANGE
        p.paragraph_format.space_after = Pt(24)

    def add_heading1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)

    def add_heading2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = GOLD
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)

    def add_heading3(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def style_table(table):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header formatting
        for cell in table.rows[0].cells:
            set_cell_background(cell, "1E293B")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Alternate row shading
        for i, row in enumerate(table.rows[1:], start=1):
            fill = "F8FAFC" if i % 2 == 1 else "FFFFFF"
            for cell in row.cells:
                set_cell_background(cell, fill)

    # ------------------------------------------------------------- Document Body

    # Title & Subtitle
    add_title("Restaurant QR Code Ordering & Management System")
    add_subtitle("Complete Technical Architecture, Credentials & User Manual")
    add_brand_subtitle("Designed & Developed by MONK DEVELOPER")

    # Metadata Box
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("System Version: 1.0.0  |  Technology: React 19 + Node.js Express 5 + PostgreSQL  |  Date: July 2026")
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = DARK_GRAY
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 1. Executive Summary
    add_heading1("1. Executive Summary & System Overview")
    doc.add_paragraph(
        "The Restaurant QR Code Ordering & Management System is an enterprise-grade, full-stack web application "
        "designed to streamline dining operations. Diners scan a unique table-specific QR code to browse the menu, "
        "place orders without account creation, and monitor order cooking status in real time via WebSockets. "
        "Simultaneously, restaurant staff, chefs, and administrators view and manage orders from their respective specialized interfaces."
    )

    add_heading2("The Five Key Portals")
    portals_table = doc.add_table(rows=1, cols=3)
    hdr_cells = portals_table.rows[0].cells
    hdr_cells[0].text = "Portal Name"
    hdr_cells[1].text = "Target Audience & URL"
    hdr_cells[2].text = "Key Functionality"

    portals_data = [
        ("Customer Ordering", "Diners (No login)\n/t/:token", "Browse digital menu, filter categories, place orders, view live order tracking."),
        ("Kitchen Display (KDS)", "Chefs & Kitchen Staff\n/kitchen", "4-column live order board (Pending -> Preparing -> Ready -> Completed) with 1-tap status updates."),
        ("Staff / Waiter Console", "Waiters & Floor Staff\n/staff", "Manage table orders, accept cash/card payments, print itemized invoices, track table status."),
        ("Admin Dashboard", "Restaurant Managers\n/admin", "Real-time revenue metrics, top-selling items, menu management, table & QR code generation."),
        ("Super Admin / Security", "System Administrator\n/admin/users", "Staff account creation, Role-Based Access Control (RBAC), security audit logs.")
    ]

    for p_name, p_url, p_desc in portals_data:
        row = portals_table.add_row().cells
        row[0].text = p_name
        row[1].text = p_url
        row[2].text = p_desc

    style_table(portals_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 2. Technology Stack
    add_heading1("2. Complete Technology Stack")
    
    doc.add_paragraph(
        "The application is built using modern, type-safe, and industry-standard open-source technologies."
    )

    add_heading2("Frontend Architecture")
    doc.add_paragraph("• Core Framework: React 19 with TypeScript (Type-safe client implementation)")
    doc.add_paragraph("• Build & Dev Tool: Vite 8 (Hot Module Replacement and fast bundling)")
    doc.add_paragraph("• Styling & Design System: Tailwind CSS v4 with custom Cormorant Garamond & Jost variable fonts")
    doc.add_paragraph("• State & Data Fetching: TanStack React Query v5 (Server-state caching, invalidation, and polling)")
    doc.add_paragraph("• HTTP Client & Routing: Axios with credentials interceptors, React Router v7")
    doc.add_paragraph("• Real-Time Engine: Socket.io-Client (WebSocket synchronization for order events)")
    doc.add_paragraph("• Schema Validation & Icons: Zod, Lucide-React Icons")

    add_heading2("Backend Architecture")
    doc.add_paragraph("• Runtime & Server: Node.js (v20+) with Express.js 5 and TypeScript (TSX runtime)")
    doc.add_paragraph("• Database & ORM: PostgreSQL with Prisma ORM v7 (19 relational tables with migrations)")
    doc.add_paragraph("• Real-Time WebSocket Server: Socket.io Server (Room-scoped broadcasting)")
    doc.add_paragraph("• Authentication & Security: JWT (JSON Web Tokens), Bcrypt (Password hashing), Helmet (Security HTTP headers), Express Rate Limit, Cookie Parser")
    doc.add_paragraph("• File Uploads & Utilities: QRCode (PNG generation), Multer (Image uploads), Compression (gzip)")
    doc.add_paragraph("• API Documentation: Swagger UI Dist, OpenAPI 3.1 Spec, Redocly CLI")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. System Credentials & Reference
    add_heading1("3. Credentials & Access Reference")

    doc.add_paragraph("Default password for all seed demo accounts: DemoPassword2026")

    cred_table = doc.add_table(rows=1, cols=4)
    c_hdrs = cred_table.rows[0].cells
    c_hdrs[0].text = "Role"
    c_hdrs[1].text = "Email"
    c_hdrs[2].text = "Password"
    c_hdrs[3].text = "Portal Route"

    creds = [
        ("Chef (Kitchen)", "chef@restaurant.local", "Env: SEED_CHEF_PASSWORD", "http://localhost:5173/kitchen"),
        ("Staff (Waiter)", "waiter@restaurant.local", "Env: SEED_WAITER_PASSWORD", "http://localhost:5173/staff"),
        ("Admin (Manager)", "manager@restaurant.local", "Env: SEED_MANAGER_PASSWORD", "http://localhost:5173/admin"),
        ("Super Admin", "admin@restaurant.local", "Env: SEED_ADMIN_PASSWORD", "http://localhost:5173/login")
    ]

    for r, e, p, u in creds:
        row = cred_table.add_row().cells
        row[0].text = r
        row[1].text = e
        row[2].text = p
        row[3].text = u

    style_table(cred_table)

    add_heading2("Database Connection Details")
    doc.add_paragraph("• Database Host: localhost")
    doc.add_paragraph("• Database Port: 5432")
    doc.add_paragraph("• Database Name: restaurant_db")
    doc.add_paragraph("• Username: <POSTGRES_USER>")
    doc.add_paragraph("• Password: <POSTGRES_PASSWORD>")
    doc.add_paragraph("• Connection String (DATABASE_URL): postgresql://<POSTGRES_USER>:<POSTGRES_PASSWORD>@localhost:5432/restaurant_db")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 4. Installation & Local Setup Guide
    add_heading1("4. Step-by-Step Installation & Setup Guide")

    add_heading2("Prerequisites")
    doc.add_paragraph("1. Node.js v20.0.0 or higher installed.")
    doc.add_paragraph("2. Docker Desktop running (for PostgreSQL container).")
    doc.add_paragraph("3. Git installed.")

    add_heading2("Step 1: Database & Backend Setup")
    doc.add_paragraph("Execute the following commands in the terminal inside the server/ directory:")
    
    code1 = doc.add_paragraph()
    code1_run = code1.add_run(
        "cd server\n"
        "cp .env.example .env\n"
        "docker compose up -d\n"
        "npm install\n"
        "npm run db:migrate\n"
        "npm run seed\n"
        "npm run seed:demo\n"
        "npm run dev"
    )
    code1_run.font.name = 'Consolas'
    code1_run.font.size = Pt(9.5)

    add_heading2("Step 2: Frontend Setup")
    doc.add_paragraph("Open a second terminal window and execute the following commands inside the client/ directory:")

    code2 = doc.add_paragraph()
    code2_run = code2.add_run(
        "cd client\n"
        "cp .env.example .env\n"
        "npm install\n"
        "npm run dev"
    )
    code2_run.font.name = 'Consolas'
    code2_run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 5. Core System Features
    add_heading1("5. Key Features & Business Logic")

    add_heading2("1. Table Reservation & Dynamic Seat Engine")
    doc.add_paragraph(
        "The reservation module (/reserve) features a dynamic seat capacity calculation engine. "
        "When a user selects a date, time slot, and party size, the system aggregates total active table capacity "
        "and subtracts all existing bookings (PENDING, CONFIRMED, SEATED) for that exact window. "
        "It dynamically displays: Total Capacity, Free Seats at selected time, Selected Guests, and Seats Remaining after booking."
    )

    add_heading2("2. Real-Time Order Synchronization")
    doc.add_paragraph(
        "Powered by Socket.io rooms. Order events (Created, Status Changed, Cancelled) are emitted instantaneously. "
        "Kitchen displays, staff consoles, and customer tracking screens update in real time without manual page refreshes."
    )

    add_heading2("3. Dynamic CORS & Mobile Local Area Network (LAN) Support")
    doc.add_paragraph(
        "The application automatically resolves local IP addresses (e.g. 10.128.48.93 or 172.22.224.1). "
        "When a customer scans a table QR code on their smartphone over Wi-Fi, the frontend dynamically points to the server's "
        "network origin, bypassing CORS limitations and localhost isolation."
    )

    add_heading2("4. Table & QR Management")
    doc.add_paragraph(
        "Admins can add new tables specifying both Table Number (e.g. T-09) and Seat Capacity (e.g. 2, 4, 6, 8 seats). "
        "The system generates a unique QR code image on demand, ready for high-resolution PNG download."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 6. API Reference
    add_heading1("6. API Reference Overview")

    api_table = doc.add_table(rows=1, cols=3)
    a_hdrs = api_table.rows[0].cells
    a_hdrs[0].text = "Method & Endpoint"
    a_hdrs[1].text = "Auth & Permission"
    a_hdrs[2].text = "Description"

    apis = [
        ("GET /api/foods", "Public", "List menu items with category & offer filters"),
        ("GET /api/categories", "Public", "List active food categories"),
        ("POST /api/orders", "Public (Customer)", "Place a new table order with items"),
        ("GET /api/orders/track/:token", "Public (Customer)", "Track live order status by token"),
        ("GET /api/reservations/availability", "Public", "Query available seats for date/time/partySize"),
        ("POST /api/reservations", "Public", "Request a table reservation"),
        ("POST /api/auth/login", "Public", "Authenticate staff and issue JWT tokens"),
        ("GET /api/tables", "Staff Auth", "List all tables with QR tokens and status"),
        ("POST /api/tables", "Admin (table:create)", "Create new table with seat capacity"),
        ("GET /api/admin/reports/dashboard", "Admin (dashboard:view)", "Get live revenue, order metrics & top sellers"),
        ("GET /api/docs", "Public", "Interactive Swagger UI documentation page")
    ]

    for m, a, d in apis:
        row = api_table.add_row().cells
        row[0].text = m
        row[1].text = a
        row[2].text = d

    style_table(api_table)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Final note
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run("--- End of Documentation ---")
    r_end.font.italic = True
    r_end.font.color.rgb = DARK_GRAY

    # Save document
    doc_path = "C:\\Users\\Ankita\\OneDrive\\Desktop\\Restaurant-QR-Ordering-System\\Restaurant-QR-Ordering-System-Documentation.docx"
    try:
        doc.save(doc_path)
        print(f"Documentation saved successfully to {doc_path}")
    except PermissionError:
        doc_path_alt = "C:\\Users\\Ankita\\OneDrive\\Desktop\\Restaurant-QR-Ordering-System\\Restaurant-QR-Ordering-System-Documentation-Updated.docx"
        doc.save(doc_path_alt)
        print(f"File locked by Word. Documentation saved successfully to {doc_path_alt}")

if __name__ == "__main__":
    create_documentation()
